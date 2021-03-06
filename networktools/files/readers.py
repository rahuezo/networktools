import os, sys

from extensions import ExtensionHandler
from sanitizer import sanitize_string


def read_docx(f, raw=False):
    """
    This function gets the alphanumeric content in a docx file.

    Args:
        f: This is the input file to be read.

    Returns: 
        An alphanumeric string representation of the contents of f.
    
    Raises:
        ImportError: Module python-docx needs to be installed first.
    """ 
    try: 
        import docx
    except ImportError: 
        print 'You need to install docx. Try, sudo pip install python-docx'
        sys.exit()

    document = docx.Document(f)
    if raw: 
        return ''.join([p.text for p in document.paragraphs])
    return ''.join([sanitize_string(p.text) for p in document.paragraphs])


def read_txt(f, raw=False):
    """
    This function gets the alphanumeric content in a txt file.

    Args:
        f: This is the input file to be read.

    Returns: 
        An alphanumeric string representation of the contents of f.
    
    Raises:
        None.
    """ 
    if raw: 
        return f.read()
    return sanitize_string(f.read())


class FileReader(ExtensionHandler): 
    """
    This class inherits from ExtensionHandler to detect file extensions
    and lets you read the contents of a file.
    """

    def __init__(self, f): 
        ExtensionHandler.__init__(self, f)

    def read(self, raw=False): 
        extension = self.get_extension()

        if extension == 'docx': 
            return read_docx(self.f, raw=raw)

        elif extension == 'txt': 
            return read_txt(self.f, raw=raw)
        else: 
            print '\nThis is an unsupported file of type {}\n'.format(extension)
            return None
